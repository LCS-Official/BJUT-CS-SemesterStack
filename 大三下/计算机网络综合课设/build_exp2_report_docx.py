from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parent
IMG = ROOT / "实验二" / "实验2截图"
OUT = ROOT / "计网课设实验二报告-LC.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run(run, bold=bold)


def set_run(run, size=12, bold=False, font="宋体"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def set_para_format(p, first_line=True, align=None, space_after=Pt(4)):
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = space_after
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align


def para(doc, text="", first_line=True, align=None, size=12, bold=False):
    p = doc.add_paragraph()
    set_para_format(p, first_line=first_line, align=align)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    return p


def heading(doc, text):
    p = para(doc, text, first_line=False, size=14, bold=True)
    p.paragraph_format.space_before = Pt(8)
    return p


def subheading(doc, text):
    p = para(doc, text, first_line=False, size=12, bold=True)
    p.paragraph_format.space_before = Pt(4)
    return p


def image(doc, filename, caption):
    path = IMG / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.4))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run(r, size=10.5, font="宋体")


def table(doc, rows):
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, value in enumerate(rows[0]):
        set_cell_text(t.rows[0].cells[i], value, bold=True)
    for row in rows[1:]:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    return t


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.7)
sec.right_margin = Cm(2.7)

styles = doc.styles
styles["Normal"].font.name = "宋体"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
styles["Normal"].font.size = Pt(12)

para(doc, "计算机网络综合课设", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True)
para(doc, "实验报告", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True)
para(doc, "", first_line=False)

cover = [
    ("专业班级", "230702班"),
    ("姓名", "王宏天"),
    ("学号", "23070219"),
    ("组长学号", "23070219"),
    ("组长姓名", "王宏天"),
    ("组长联系方式", "13011884082"),
    ("组员学号、姓名", "23071005 侯景祺；23070225 贺子淇；23070226 尹先炜；23070214 宋天翔"),
    ("实验日期", "2026/7/8"),
    ("实验名称", "实验二 数据包的捕获与分析"),
]

ct = doc.add_table(rows=len(cover), cols=2)
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
ct.style = "Table Grid"
for i, (k, v) in enumerate(cover):
    set_cell_text(ct.rows[i].cells[0], k, bold=True)
    set_cell_text(ct.rows[i].cells[1], v)

doc.add_page_break()

para(doc, "目录", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
for item in [
    "一、实验内容和要求",
    "二、实验环境",
    "三、实验需求分析与逻辑框图",
    "四、核心功能的实现机制",
    "五、程序源代码（核心部分）",
    "六、程序扩展功能的需求分析与实现",
    "七、实验数据、结果分析",
    "八、总结",
    "九、同组人分工情况",
]:
    para(doc, item, first_line=False)

doc.add_page_break()

heading(doc, "一、实验内容和要求")
para(doc, "本实验为“实验二 数据包的捕获与分析”。实验目的在于通过 Wireshark 软件监控局域网通信过程，捕获并保存网络数据包，利用过滤器定位特定数据流，并对常用协议的数据包格式和字段含义进行分析。")
for line in [
    "（1）安装并启动 Wireshark 软件，选择合适的网络接口进行抓包。",
    "（2）设置网卡为混杂模式，使 Wireshark 能够监控局域网通信状态。",
    "（3）启动数据包捕获，跟踪主机与网关、DNS 服务器及外部地址之间的通信报文，并保存抓包文件以便复查。",
    "（4）设置显示过滤器，分别观察 ARP、ICMP、TCP、UDP/DNS、IEEE 802.3 LLC 等典型协议或数据流。",
    "（5）对 Ethernet、IEEE 802.3、IP、ICMP、ARP、TCP、UDP 等常用协议的数据包字段进行分析，并利用统计工具查看协议层次和流量分布。",
]:
    para(doc, line)

heading(doc, "二、实验环境")
table(
    doc,
    [
        ["项目", "配置"],
        ["操作系统", "Windows 11；实验指导书参考环境为 Windows/Windows XP 终端"],
        ["协议分析软件", "Wireshark Portable 4.6.6"],
        ["抓包接口", "WLAN，接口描述为 Intel(R) Wi-Fi 7 BE201 320MHz"],
        ["本机 IPv4", "192.168.138.120/24"],
        ["网关与 DNS", "192.168.138.246"],
        ["原始抓包文件", "exp2_mix.pcapng，共 122 个包，抓包时长 23.455 s，0 丢包"],
        ["补充分析文件", "IEEE 802.3 LLC 补充分析抓包文件，用于 IEEE 802.3 LLC 帧格式说明"],
    ],
)

heading(doc, "三、实验需求分析与逻辑框图")
para(doc, "实验过程围绕“抓包—过滤—分析—统计”展开。首先在 Wireshark 中选择 WLAN 接口并开启抓包；随后通过 ping、DNS 查询和浏览器/HTTP 访问产生 ICMP、ARP、UDP、TCP 等流量；抓包完成后保存为 pcapng 文件；最后利用显示过滤器与统计功能分析各协议字段。")
para(doc, "逻辑流程如下：")
for line in [
    "选择抓包接口 → 启动混杂模式 → 开始实时捕获 → 产生测试流量 → 保存 pcapng 文件",
    "打开抓包文件 → 设置显示过滤器 → 定位典型协议数据包 → 查看协议树和字节流",
    "打开 Statistics / Protocol Hierarchy → 统计各协议包数量、字节数和占比 → 汇总分析结果",
]:
    para(doc, line)

heading(doc, "四、核心功能的实现机制")
subheading(doc, "1、Wireshark 抓包与接口选择")
para(doc, "Wireshark 底层通过 Npcap/Dumpcap 访问网卡，捕获经过网络接口的数据帧。本次实验选择 WLAN 接口进行抓包，抓包文件封装类型为 Ethernet，抓包时开启混杂模式，未使用禁止混杂模式的 -p 参数。")
image(doc, "wireshark_gui_01_overview.png", "图 1 Wireshark 原始抓包结果总览")

subheading(doc, "2、ICMP 报文分析")
para(doc, "ICMP 报文由 ping 命令产生，主要用于测试网络连通性。包 13 为 Echo request，从 192.168.138.120 发往 192.168.138.246；包 14 为 Echo reply，说明网关能够正常响应。本实验中 ICMP 请求和响应的 Identifier 与 Sequence Number 对应一致，可用于匹配一次 ping 过程。")
image(doc, "wireshark_gui_02_icmp_filter.png", "图 2 ICMP 报文过滤与请求/响应结果")

subheading(doc, "3、ARP 报文分析")
para(doc, "ARP 用于在局域网内根据 IPv4 地址解析目标主机的 MAC 地址。实验中可以观察到主机向广播地址发送 Who has 查询，并接收网关返回的 is at 响应。ARP 报文字段中包括硬件类型、协议类型、发送方 MAC、发送方 IP、目标 MAC 和目标 IP，是 IP 通信前完成二层寻址的重要过程。")
image(doc, "wireshark_gui_03_arp_filter.png", "图 3 ARP 报文过滤结果")

subheading(doc, "4、TCP 报文分析")
para(doc, "TCP 是面向连接的传输层协议。实验中包 24 为本机 192.168.138.120 向外部主机 199.59.149.203 的 443 端口发送 SYN 报文，表示连接建立请求。TCP 头部中可以看到源端口、目的端口、序号、确认号、窗口大小、标志位和选项字段。后续重传报文说明目标连接未能立即完成建立。")
image(doc, "wireshark_gui_04_tcp_ipv4_filter.png", "图 4 TCP/IPv4 报文过滤结果")

subheading(doc, "5、UDP/DNS 报文分析")
para(doc, "UDP 是无连接传输层协议，DNS 查询通常基于 UDP 53 端口完成。实验中 DNS 查询包显示本机向 192.168.138.246 发送 example.com 的 A 记录查询，随后收到包含解析结果的 DNS response。UDP 报文头部字段较少，主要包括源端口、目的端口、长度和校验和。")
image(doc, "wireshark_gui_05_udp_dns_filter.png", "图 5 UDP/DNS 报文过滤结果")

subheading(doc, "6、IEEE 802.3 LLC 帧分析")
para(doc, "IEEE 802.3 帧与 Ethernet II 帧的重要区别在于第 13-14 字节字段含义不同：Ethernet II 中该字段表示上层协议类型；IEEE 802.3 中该字段表示数据长度，且其数值小于 0x0600。本次补充分析的 LLC 帧中 Length 字段为 46，后续为 Logical-Link Control 头部，包含 DSAP、SSAP 和 Control 字段。")
para(doc, "Wireshark 协议树中可见 IEEE 802.3 Ethernet 和 Logical-Link Control 两层，说明该帧按照 IEEE 802.3 LLC 格式解析。")
image(doc, "wireshark_gui_07_ieee8023_llc_demo.png", "图 6 IEEE 802.3 LLC 帧分析结果")

subheading(doc, "7、协议层次统计")
para(doc, "Wireshark 的 Protocol Hierarchy Statistics 可以按协议层次统计包数量、字节数和占比。本次原始抓包共 122 个包，包含 Ethernet、IPv4、IPv6、TCP、UDP、DNS、ICMP、ARP、802.1Q VLAN 等协议，可用于从整体上观察网络流量组成。")
image(doc, "wireshark_gui_06_protocol_hierarchy.png", "图 7 Wireshark 协议层次统计结果")

heading(doc, "五、程序源代码（核心部分）")
para(doc, "本实验属于协议分析实验，不涉及自编程序源码。核心操作为 Wireshark 抓包、保存文件和显示过滤器配置。关键过滤器如下表所示。")
table(
    doc,
    [
        ["分析对象", "显示过滤器", "典型包号/说明"],
        ["ARP", "arp", "21/22/39/40，地址解析请求与响应"],
        ["ICMP", "icmp", "13/14，Echo request/reply"],
        ["TCP over IPv4", "tcp and ip", "24，SYN 连接请求"],
        ["UDP/DNS", "dns or udp", "80/82，DNS 查询与响应"],
        ["IEEE 802.3 LLC", "llc", "1，IEEE 802.3 Ethernet + Logical-Link Control"],
    ],
)
para(doc, "抓包保存文件：exp2_mix.pcapng；按协议导出的过滤文件包括 filter_arp.pcapng、filter_icmp.pcapng、filter_tcp_ipv4.pcapng 和 filter_udp_dns.pcapng。")

heading(doc, "六、程序扩展功能的需求分析与实现")
para(doc, "本实验没有程序扩展开发任务。为增强实验报告的完整性，除基本抓包和过滤外，还进行了以下补充分析：")
for line in [
    "（1）按协议分别导出 ARP、ICMP、TCP/IPv4、UDP/DNS 抓包文件，便于单独复查。",
    "（2）使用 Protocol Hierarchy Statistics 统计协议层次、包数量和字节数。",
    "（3）对 IEEE 802.3 LLC 帧单独进行补充分析，比较其 Length 字段与 Ethernet II Type 字段的区别。",
]:
    para(doc, line)

heading(doc, "七、实验数据、结果分析")
subheading(doc, "1、抓包总体结果")
table(
    doc,
    [
        ["指标", "结果"],
        ["原始抓包文件", "exp2_mix.pcapng"],
        ["抓包接口", "WLAN"],
        ["包数量", "122"],
        ["抓包时长", "23.455320800 s"],
        ["数据大小", "47 kB"],
        ["平均包大小", "388.11 bytes"],
        ["丢包情况", "0 丢包"],
    ],
)
subheading(doc, "2、协议分析结果")
for line in [
    "（1）Ethernet：所有数据包均以 Ethernet 封装进入 Wireshark，数据链路层可观察源 MAC、目的 MAC 和类型/长度字段。",
    "（2）ARP：主机通过广播查询目标 IP 对应的 MAC 地址，并通过单播响应完成地址解析。",
    "（3）IP：IPv4 报文包含版本号、首部长度、总长度、TTL、协议号、源地址和目的地址等字段，是上层 ICMP/TCP/UDP 的承载层。",
    "（4）ICMP：Echo request 与 Echo reply 成对出现，用于测试本机与网关之间的连通性，TTL 和序列号字段可用于判断报文路径和匹配关系。",
    "（5）TCP：SYN 报文体现 TCP 建立连接的第一步，源端口为临时端口，目的端口为服务端口，标志位 SYN 置 1。",
    "（6）UDP/DNS：DNS 查询和响应体现 UDP 无连接通信特点，报文中包含查询名称、查询类型、响应地址等字段。",
    "（7）IEEE 802.3 LLC：Length 字段取代 EtherType 字段，后接 LLC 头部，DSAP/SSAP 表示服务访问点，Control 表示 LLC 控制信息。",
]:
    para(doc, line)

heading(doc, "八、总结")
para(doc, "通过本次实验，掌握了 Wireshark 的基本安装、启动、接口选择、混杂模式抓包、显示过滤器设置和协议统计功能。实验过程中对 ARP、IP、ICMP、TCP、UDP/DNS、Ethernet II 等协议进行了逐项观察，并结合 IEEE 802.3 LLC 补充分析样例识别数据链路层帧结构差异。")
para(doc, "本次抓包说明，协议分析不仅可以验证网络连通性，还可以帮助定位 DNS 解析、TCP 连接建立和局域网地址解析等具体过程。通过 Protocol Hierarchy Statistics 可以快速了解网络流量组成，为后续网络管理、故障排查和安全分析奠定基础。")

heading(doc, "九、同组人分工情况")
table(
    doc,
    [
        ["学号", "姓名", "承担任务"],
        ["23070219", "王宏天", "Wireshark 抓包、协议字段分析、报告整理"],
        ["23071005", "侯景祺", "实验环境检查、截图整理"],
        ["23070225", "贺子淇", "ARP/ICMP 报文分析"],
        ["23070226", "尹先炜", "TCP/UDP/DNS 报文分析"],
        ["23070214", "宋天翔", "统计结果整理与复核"],
    ],
)

doc.save(OUT)
print(OUT)
