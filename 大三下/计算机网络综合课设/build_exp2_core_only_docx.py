from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parent
IMG = ROOT / "实验二" / "实验2截图"
OUT = ROOT / "计网课设实验二报告-LC.docx"
FALLBACK_OUT = ROOT / "计网课设实验二报告-LC-核心功能.docx"


def set_run(run, size=12, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    rpr = run._element.get_or_add_RPr() if hasattr(run._element, "get_or_add_RPr") else run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")


def para(doc, text="", first_line=True, size=12, bold=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)
    return p


def heading(doc, text, size=14):
    p = para(doc, text, first_line=False, size=size, bold=True)
    p.paragraph_format.space_before = Pt(8)
    return p


def image(doc, filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(IMG / filename), width=Inches(6.4))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    set_run(r, size=10.5)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.7)
sec.right_margin = Cm(2.7)

doc.styles["Normal"].font.name = "宋体"
doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
doc.styles["Normal"].font.size = Pt(12)

heading(doc, "四、核心功能的实现机制")
heading(doc, "2、实验二", size=12)

para(doc, "实验二的核心功能包括数据包捕获、显示过滤、协议字段解析和统计分析四部分。Wireshark 通过底层抓包驱动读取网卡收到和发出的链路层数据帧，并将其保存为 pcapng 文件。实验中选择 WLAN 网卡作为抓包接口，启动实时捕获后产生 ping、DNS 查询和 TCP 连接等测试流量，从而获得 ARP、ICMP、TCP、UDP/DNS 等典型报文。")
para(doc, "过滤机制分为捕获前过滤和捕获后过滤。本实验主要使用捕获后显示过滤器，在已经保存的抓包文件中通过 arp、icmp、tcp and ip、dns or udp、llc 等过滤条件定位目标协议。显示过滤不会删除原始数据，只改变当前窗口中显示的数据包，便于对同一份抓包文件反复分析。")
para(doc, "协议解析机制体现为 Wireshark 的分层协议树。一个数据包从 Frame 开始，逐层解析 Ethernet、IP、TCP/UDP/ICMP 等协议字段。对于 Ethernet II 帧，第 13-14 字节解释为类型字段；对于 IEEE 802.3 帧，该字段解释为长度字段，后续再由 LLC 头部中的 DSAP、SSAP、Control 字段说明链路层服务访问点和控制信息。")
para(doc, "统计分析机制通过 Statistics 菜单中的 Protocol Hierarchy Statistics 实现。该工具按协议层次统计数据包数量、字节数和占比，可以从整体上判断本次抓包中各协议流量的组成情况，为后续结果分析提供依据。")

heading(doc, "七、实验数据、结果分析")
heading(doc, "2、实验二", size=12)

para(doc, "本次实验原始抓包文件共捕获 122 个数据包，抓包时长约 23.455 s，抓包接口为 WLAN，文件封装类型为 Ethernet。抓包过程中通过不同显示过滤器对典型协议进行定位，并结合协议树字段解释各报文的作用。")
image(doc, "wireshark_gui_01_overview.png", "图 1 Wireshark 原始抓包结果总览")

heading(doc, "（1）ICMP 报文结果分析", size=12)
para(doc, "使用过滤器 icmp 后，可以看到 ICMP Echo request 与 Echo reply 成对出现。包 13 为本机 192.168.138.120 发往 192.168.138.246 的 Echo request，包 14 为对端返回的 Echo reply。两者的 Identifier 和 Sequence Number 对应一致，说明本机与网关之间的连通性正常。")
image(doc, "wireshark_gui_02_icmp_filter.png", "图 2 ICMP 报文过滤与请求/响应结果")

heading(doc, "（2）ARP 报文结果分析", size=12)
para(doc, "使用过滤器 arp 后，可以观察到 ARP 请求与响应。主机通过广播形式发送 Who has 查询，请求目标 IP 对应的 MAC 地址；目标设备返回 is at 响应，给出对应的硬件地址。该结果说明 IP 报文发送前需要先完成局域网内的二层地址解析。")
image(doc, "wireshark_gui_03_arp_filter.png", "图 3 ARP 报文过滤结果")

heading(doc, "（3）TCP 报文结果分析", size=12)
para(doc, "使用过滤器 tcp and ip 后，可以看到 IPv4 上的 TCP 报文。包 24 为本机向 199.59.149.203 的 443 端口发送的 SYN 报文，表示 TCP 连接建立请求。协议树中可见源端口、目的端口、序号、窗口大小和 SYN 标志位等字段，后续重传报文说明该连接未能立即完成建立。")
image(doc, "wireshark_gui_04_tcp_ipv4_filter.png", "图 4 TCP/IPv4 报文过滤结果")

heading(doc, "（4）UDP/DNS 报文结果分析", size=12)
para(doc, "使用过滤器 dns or udp 后，可以观察到 DNS 查询和响应。DNS 查询报文由本机发往 DNS 服务器，响应报文返回域名对应的地址记录。UDP 报文头部字段包括源端口、目的端口、长度和校验和，结构较 TCP 简单，不需要连接建立过程。")
image(doc, "wireshark_gui_05_udp_dns_filter.png", "图 5 UDP/DNS 报文过滤结果")

heading(doc, "（5）IEEE 802.3 LLC 帧结果分析", size=12)
para(doc, "使用过滤器 llc 后，可以看到 IEEE 802.3 Ethernet 与 Logical-Link Control 两层。该帧中 Length 字段为长度含义，而不是 Ethernet II 的 EtherType 类型含义；LLC 头部继续给出 DSAP、SSAP 和 Control 字段，用于说明链路层服务访问点和控制方式。")
image(doc, "wireshark_gui_07_ieee8023_llc_demo.png", "图 6 IEEE 802.3 LLC 帧分析结果")

heading(doc, "（6）协议层次统计结果分析", size=12)
para(doc, "Protocol Hierarchy Statistics 显示，本次抓包包含 Ethernet、IPv4、IPv6、TCP、UDP、DNS、ICMP、ARP、802.1Q VLAN 等协议。其中 IPv4 下包含 TCP、UDP、ICMP 等常见协议，UDP 下包含 DNS 查询流量，ARP 用于局域网地址解析。该统计结果从整体上反映了抓包文件中的协议组成。")
image(doc, "wireshark_gui_06_protocol_hierarchy.png", "图 7 Wireshark 协议层次统计结果")

try:
    doc.save(OUT)
    print(OUT)
except PermissionError:
    doc.save(FALLBACK_OUT)
    print(FALLBACK_OUT)
