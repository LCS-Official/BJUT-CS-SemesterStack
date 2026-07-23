from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "文档与绘图"
TODAY = "2026-07-11"
PROJECT = "智能实验室预约与设备管理系统"
GROUP = "230710_8"
AUTHOR = "项目组（B 模块已按统一入口同步）"
HEADING_FONT = "微软雅黑"


OVERVIEW_OUT = OUT_DIR / f"概要设计说明书_初稿_{TODAY}.docx"
DETAIL_OUT = OUT_DIR / f"详细设计说明书_初稿_{TODAY}.docx"
ARTIFACT_OUT = OUT_DIR / "_design_docs_artifact.md"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, align: str = "left") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(str(text))
    run.bold = bold
    set_run_font(run, size=10.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(
    run,
    *,
    size: float = 10.5,
    bold: bool | None = None,
    color: str | None = None,
    ascii_font: str = "Times New Roman",
    east_asia_font: str = "宋体",
) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:cs"), ascii_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 14, "1F4E79"),
        ("Heading 3", 12, "333333"),
    ]:
        style = doc.styles[name]
        style.font.name = HEADING_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), HEADING_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), HEADING_FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = PROJECT
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    add_field(footer, "PAGE")
    footer.add_run(" 页")
    for run in footer.runs:
        set_run_font(run, size=9, color="666666")


def add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instr} "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_title_page(doc: Document, doc_name: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run(PROJECT)
    set_run_font(run, size=22, bold=True, ascii_font=HEADING_FONT, east_asia_font=HEADING_FONT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(doc_name)
    set_run_font(run, size=20, bold=True, color="1F4E79", ascii_font=HEADING_FONT, east_asia_font=HEADING_FONT)

    meta = [
        ("课程", "计算机软件综合课设"),
        ("小组", GROUP),
        ("版本", f"初稿（{TODAY}）"),
        ("编写说明", "供组长整合；B 资源与预约模块已按统一 Flask 主入口同步"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (k, v) in enumerate(meta):
        set_cell_text(table.cell(i, 0), k, bold=True, align="center")
        set_cell_text(table.cell(i, 1), v)
        set_cell_shading(table.cell(i, 0), "D9EAF7")
    set_table_widths(table, [2500, 6500])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run(f"完成日期：{TODAY}")
    set_run_font(run, size=12)
    doc.add_section(WD_SECTION.NEW_PAGE)


def normalize_heading_runs(doc: Document) -> None:
    """Force every heading run to one font to avoid Word/WPS fallback mixing."""

    sizes = {"Heading 1": 16, "Heading 2": 14, "Heading 3": 12}
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        if style_name not in sizes:
            continue
        for run in paragraph.runs:
            set_run_font(
                run,
                size=sizes[style_name],
                bold=True,
                color="1F4E79" if style_name in {"Heading 1", "Heading 2"} else "333333",
                ascii_font=HEADING_FONT,
                east_asia_font=HEADING_FONT,
            )


def add_toc_placeholder(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    p = doc.add_paragraph("提示：在 Word/WPS 中打开后，可右键更新目录。")
    p.paragraph_format.first_line_indent = None
    add_toc_field(doc.add_paragraph())
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r' TOC \o "1-3" \h \z \u '
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在打开文档后更新"
    fld_char_3 = OxmlElement("w:fldChar")
    fld_char_3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    run._r.append(text)
    run._r.append(fld_char_3)


def add_para(doc: Document, text: str = "", *, style: str | None = None, indent: bool = True) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    run = p.add_run(text)
    set_run_font(run, size=10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        set_cell_text(table.cell(0, j), h, bold=True, align="center")
        set_cell_shading(table.cell(0, j), "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            align = "center" if len(str(value)) <= 12 else "left"
            set_cell_text(cells[j], value, align=align)
    if widths:
        set_table_widths(table, widths)
    doc.add_paragraph()


def set_table_widths(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def build_overview() -> None:
    doc = Document()
    apply_styles(doc)
    add_title_page(doc, "概要设计说明书")
    add_toc_placeholder(doc)

    doc.add_heading("1．引言", level=1)
    doc.add_heading("1.1 编写目的", level=2)
    add_para(doc, "本文档依据《需求规格说明书0706》、组内 Handoff、A 统一入口交接材料以及 B 模块当前代码编写，用于说明智能实验室预约与设备管理系统的总体设计、模块划分、接口关系、数据结构和运行策略。本文档供指导教师、项目组成员、编码实现人员和测试人员使用。")
    doc.add_heading("1.2 项目背景", level=2)
    add_para(doc, "高校实验室存在设备使用高峰集中、线下预约冲突、故障处理不及时、信用约束缺少数据依据等问题。本项目面向学生、实验室管理员和系统管理员，建设覆盖资源查询、预约审批、现场签到、设备报修、信用管理、通知日志和统计分析的在线管理系统。")
    doc.add_heading("1.3 定义", level=2)
    add_table(doc, ["术语", "说明"], [
        ["RBAC", "基于角色的访问控制，用户、角色、权限分离，支持一个用户绑定多个角色。"],
        ["预约单", "学生针对实验室或设备提交的使用申请记录。"],
        ["幂等号", "一次业务提交的唯一标识，本系统在预约提交中使用 HTTP Idempotency-Key。"],
        ["IoT 模拟", "课程范围内用后端接口和日志模拟门禁或设备供电指令。"],
        ["统一入口", "全组在线演示使用的 Flask app factory：unified_app:create_app()。"],
    ], [1800, 7600])
    doc.add_heading("1.4 参考资料", level=2)
    add_bullets(doc, [
        "《需求规格说明书0706.docx》",
        "《组内分工与handoff_并行软件版.md》",
        "《A_统一Flask主入口与线上部署交接文档_2026-07-10.md》",
        "《B_资源与预约模块统一入口交接文档_2026-07-11.md》",
        "LC_proj/b_reservation 当前代码与测试用例",
        "软件工程文档模板 04_概要设计说明书.dot",
    ])

    doc.add_heading("2．任务概述", level=1)
    doc.add_heading("2.1 目标", level=2)
    add_para(doc, "系统总体目标是建立覆盖“资源预约、现场使用、设备维护、信用约束、数据统计”的智能实验室管理平台，使实验室资源使用过程可预约、可审批、可追踪、可追责、可统计。")
    add_bullets(doc, [
        "学生能够查询实验室和设备可用时段，并提交预约申请。",
        "实验室管理员能够审批预约、处理报修、维护设备状态和可用时段。",
        "系统能够拦截时间冲突、故障设备预约、信用不足预约等非法申请。",
        "现场使用通过签到签退、二维码/定位校验和 IoT 模拟日志形成闭环。",
        "信用扣分、封禁、恢复、通知和操作日志保证业务可追踪。",
    ])
    doc.add_heading("2.2 运行环境", level=2)
    add_table(doc, ["项目", "设计基准"], [
        ["部署形态", "单服务器模块化单体 Web 系统，浏览器通过 HTTP/HTTPS 访问。"],
        ["语言环境", "Python 3.12 及以上；本地开发使用 conda csv 环境，服务器使用 venv。"],
        ["Web 框架", "Flask 3.1.3，线上通过 Gunicorn + Nginx 提供服务。"],
        ["数据库", "SQLite 3.45+，单一数据库文件由服务器持有，路径通过 DATABASE_PATH 注入。"],
        ["时间与编码", "Asia/Shanghai；接口和数据库使用 YYYY-MM-DDTHH:MM:SS；UTF-8 与 JSON。"],
    ], [2200, 7200])
    doc.add_heading("2.3 需求概述", level=2)
    add_para(doc, "系统按业务职责划分为用户权限、资源查询与预约、审批与现场使用、设备与报修、信用与封禁、通知与日志、统计分析、系统参数八类功能。其中当前代码已完成并验证 A/B/C 主入口集成，D 信用与运营模块预留接入位置。")
    doc.add_heading("2.4 条件与限制", level=2)
    add_bullets(doc, [
        "课程周期内优先完成核心闭环，IoT 设备以模拟接口和日志记录实现。",
        "浏览器定位存在误差，只作为辅助校验，不作为唯一处罚依据。",
        "实验室管理员只能管理自己管辖实验室内的预约、设备和报修。",
        "SQLite 适合课程规模的单机部署；后续如迁移数据库，应保持业务状态和接口语义不变。",
        "D 模块未接入前，信用/封禁资格可由 A 提供临时适配，最终应替换为 D 的真实判断。"
    ])

    doc.add_heading("3．总体设计", level=1)
    doc.add_heading("3.1 处理流程", level=2)
    add_numbered(doc, [
        "用户登录后，A 平台模块从服务端 session 恢复当前用户、当前角色、权限集合和管辖范围。",
        "学生查询资源时，C 模块根据实验室、设备状态、可用时段和已有预约返回可预约资源。",
        "学生提交预约时，B 模块读取 A 的当前用户、D 的信用资格、C 的资源可约性，并在事务内完成冲突检测和预约写入。",
        "实验室管理员审批时，B 通过 A 校验管辖范围，再次检查资源可用性和冲突，更新预约状态并记录审计日志。",
        "学生现场签到签退时，C 校验预约、二维码、定位和时间窗口，写入签到记录，并调用 B 更新预约状态。",
        "设备报修或维修时，C 更新设备状态，查询 B 中受影响预约并执行维护取消，通知交由 D 模块处理。",
        "D 模块根据预约、签到、取消、报修记录识别违规，生成信用日志、封禁记录和通知。",
    ])
    doc.add_heading("3.2 总体结构和模块外部设计", level=2)
    add_table(doc, ["模块", "主要职责", "对外能力"], [
        ["A 平台与集成", "登录、角色切换、RBAC、管辖范围、系统参数、操作日志、统一入口", "CURRENT_ACTOR、CAN_MANAGE_LAB、GET_PARAMETER、AUDIT、/api/auth/*"],
        ["B 资源与预约", "预约创建、幂等、冲突检测、审批、驳回、取消、预约状态联动", "register_routes、ReservationStore、/api/reservations*"],
        ["C 现场与设备", "实验室/设备、可用时段、签到签退、IoT 指令、报修维修", "RESOURCE_IS_BOOKABLE、FieldEquipmentStore、/api/equipment*、/api/check-*"],
        ["D 信用与运营", "违规识别、信用扣分、封禁、信用恢复、通知、统计", "STUDENT_IS_ELIGIBLE、通知/信用/统计接口"],
    ], [1900, 4300, 3200])
    add_para(doc, "统一入口 `unified_app:create_app()` 创建一个 Flask 应用实例，读取统一数据库路径，初始化 A/B/C 的 Store，对同一 SQLite 数据库执行各模块 schema 初始化，并把 A/C/D 提供的服务端能力以回调形式注入 B。")
    doc.add_heading("3.3 功能分配", level=2)
    add_table(doc, ["需求编号", "功能", "主要模块", "说明"], [
        ["FR-01", "登录、角色切换与权限控制", "A", "RBAC、session、管辖范围和越权审计。"],
        ["FR-02", "资源查询", "C + B", "C 判断设备状态与可用时段，B 提供已有预约占用信息。"],
        ["FR-03", "提交预约与并发冲突", "B", "幂等号、事务、半开区间冲突检测。"],
        ["FR-04", "预约审批", "B + A + C", "A 校验管辖范围，C 校验资源可约性。"],
        ["FR-05", "取消预约", "B + D", "B 更新预约状态，D 执行信用处理和通知。"],
        ["FR-06", "签到签退与 IoT", "C + B", "C 写签到记录，B 更新预约状态。"],
        ["FR-07/08", "报修维修与可用时段", "C + B", "C 管理设备状态，B 处理受影响预约。"],
        ["FR-09/10", "违规、信用与恢复", "D", "根据业务记录扣分、封禁和恢复。"],
        ["FR-11/12/13", "基础数据、通知日志、统计", "A/C/D", "参数、日志、通知和统计报表。"],
    ], [1300, 2600, 1800, 3700])

    doc.add_heading("4．接口设计", level=1)
    doc.add_heading("4.1 外部接口", level=2)
    add_table(doc, ["接口", "输入", "输出/状态"], [
        ["POST /api/auth/login", "username、password", "当前用户、角色列表、权限、管辖范围；200/401/403"],
        ["GET /api/auth/me", "服务端 session", "当前用户上下文；200/401"],
        ["POST /api/reservations", "Idempotency-Key；lab_id、equipment_id、start_time、end_time、purpose", "reservation；201/400/403/409/503"],
        ["GET /api/reservations/me", "服务端 session", "当前用户预约列表；200/401/503"],
        ["GET /api/labs/{lab_id}/reservations/pending", "lab_id；服务端校验管辖范围", "待审批列表；200/403"],
        ["POST /api/reservations/{id}/approve|reject", "comment；当前管理员来自 session", "更新后的 reservation；200/403/409"],
        ["POST /api/reservations/{id}/cancel", "maintenance_cancel 可选", "更新后的 reservation，含 late_cancel 和 credit_deduction_required"],
        ["GET/PUT /api/equipment/*", "设备、状态、可用时段参数", "设备和时段规则对象"],
        ["POST /api/check-in /api/check-out", "预约编号、二维码令牌、定位/人工核验信息", "签到记录、预约状态、IoT 结果"],
    ], [2500, 4200, 3000])
    doc.add_heading("4.2 内部接口", level=2)
    add_table(doc, ["接口/方法", "提供方", "使用方", "设计说明"], [
        ["CURRENT_ACTOR()", "A", "B/C/D", "从 Flask session 恢复当前登录用户和当前角色。"],
        ["CAN_MANAGE_LAB(user_id, lab_id)", "A", "B/C/D", "校验实验室管理员管辖范围，系统管理员默认全局可管理。"],
        ["GET_PARAMETER(key, default)", "A", "B/C/D", "读取取消免责时限、签到窗口、信用阈值等系统参数。"],
        ["AUDIT(...)", "A", "B/C/D", "记录登录、审批、参数变更、设备状态变更等关键操作。"],
        ["RESOURCE_IS_BOOKABLE(...)", "C", "B", "判断设备/实验室在目标时段是否可预约。"],
        ["STUDENT_IS_ELIGIBLE(user_id)", "D/A 临时适配", "B", "判断学生信用/封禁状态是否允许预约。"],
        ["ReservationStore.mark_using/mark_completed", "B", "C", "签到签退后更新预约状态。"],
        ["ReservationStore.list_affected_by_equipment", "B", "C", "设备维护时查询受影响预约。"],
    ], [2700, 1500, 1200, 4000])

    doc.add_heading("5．数据结构设计", level=1)
    doc.add_heading("5.1 逻辑结构设计", level=2)
    add_para(doc, "系统采用关系型数据模型，围绕用户、角色、实验室、设备、预约、签到、报修、违规、信用、通知、日志和系统参数建立核心表。用户与角色、角色与权限采用多对多关联；预约与用户、实验室、设备关联；签到、违规、信用记录通过预约或用户形成追踪链。")
    doc.add_heading("5.2 物理结构设计", level=2)
    add_table(doc, ["表", "归属", "关键字段", "说明"], [
        ["user_account / role / permission", "A", "user_id、role_code、permission_code", "账号、角色和权限基础数据。"],
        ["admin_lab_scope", "A", "user_id、lab_id", "管理员管辖范围。"],
        ["laboratory / equipment / equipment_availability", "C", "lab_id、equipment_id、rule_type、start_time、end_time", "实验室、设备和可用时段规则。"],
        ["reservation", "B", "reservation_id、request_id、user_id、lab_id、equipment_id、start_time、end_time、status", "预约、审批和取消记录；request_id 为幂等键。"],
        ["check_record / iot_command_log", "C", "reservation_id、verify_method、command_type、created_at、executed_at", "现场签到签退和 IoT 模拟日志。"],
        ["repair_report", "C", "repair_id、equipment_id、reporter_user_id、reported_at", "设备报修和维修处理记录。"],
        ["violation_record / credit_log / ban_record", "D", "violation_id、credit_log_id、ban_id、changed_at", "违规、信用变动和封禁。"],
        ["notification / operation_log", "D/A", "receiver_user_id、operator_user_id、target_object_id、operation_time", "通知和审计记录。"],
        ["system_parameter / history", "A", "param_key、version、updated_by_user_id、changed_at", "规则配置和参数变更历史。"],
    ], [2200, 900, 3600, 3000])
    doc.add_heading("5.3 数据结构与程序的关系", level=2)
    add_table(doc, ["程序模块", "主要读写表", "一致性要求"], [
        ["PlatformStore", "user_account、role、user_role、permission、admin_lab_scope、system_parameter、operation_log", "权限、参数和日志统一由 A 提供。"],
        ["ReservationStore", "reservation", "创建、审批、取消均在短事务内完成；冲突检测与写入不可分离。"],
        ["FieldEquipmentStore", "laboratory、equipment、equipment_availability、check_record、repair_report、iot_command_log", "设备状态变更应影响预约可用性。"],
        ["Credit/Operation Store", "violation_record、credit_log、credit_restore_request、ban_record、notification", "扣分、恢复、封禁和通知需幂等且可追溯。"],
    ], [2400, 4200, 3000])

    doc.add_heading("6．运行设计", level=1)
    doc.add_heading("6.1 运行模块的组合", level=2)
    add_para(doc, "线上运行以 `unified_app:create_app()` 为唯一 Flask app factory。启动时依次创建 PlatformStore、ReservationStore、FieldEquipmentStore，执行各模块 schema 初始化，注册 A/B/C 路由，并在 D 模块完成后继续注册信用与运营路由。")
    doc.add_heading("6.2 运行控制", level=2)
    add_bullets(doc, [
        "Gunicorn 以 1 worker、4 threads 方式监听 127.0.0.1:8000，由 Nginx 对外代理。",
        "业务接口统一返回 JSON，成功与失败通过 HTTP 状态码区分，错误体至少包含 error 字段。",
        "预约创建使用 Idempotency-Key 防重复提交，SQLite 使用 BEGIN IMMEDIATE 防并发穿透。",
        "关键操作调用 AUDIT 记录操作人、角色、动作、操作对象和操作时间。",
    ])
    doc.add_heading("6.3 运行时间", level=2)
    add_para(doc, "主要性能目标为：登录、资源查询 95% 请求 2 秒内返回；预约无冲突情况下 95% 在 1.5 秒内完成事务；审批 95% 在 1 秒内完成；普通统计看板 95% 在 5 秒内返回；通知在业务提交后 5 秒内生成。")

    doc.add_heading("7．出错处理设计", level=1)
    doc.add_heading("7.1 出错输出信息", level=2)
    add_table(doc, ["错误类型", "HTTP 状态", "输出"], [
        ["参数缺失或格式错误", "400", '{"error":"缺少字段或时间格式非法"}'],
        ["未登录或身份失效", "401/403", '{"error":"请先登录或无权操作"}'],
        ["预约冲突或状态非法", "409", '{"error":"预约时段冲突或当前状态不可操作"}'],
        ["必要模块未接入", "503", '{"error":"尚未接入服务端能力：..."}'],
        ["数据库或未知异常", "500", '{"error":"系统繁忙，请稍后重试"}'],
    ], [2500, 1600, 5000])
    doc.add_heading("7.2 出错处理对策", level=2)
    add_bullets(doc, [
        "数据库异常时回滚事务，避免半完成状态。",
        "IoT 指令失败不回滚已成功签到，写 failed 日志并提示管理员人工处理。",
        "定位失败允许人工核验，不直接判定违规。",
        "定时任务失败记录异常日志，允许管理员手动触发补偿扫描。",
        "网络重试依赖幂等键避免重复预约、重复扣分和重复通知。"
    ])

    doc.add_heading("8．安全保密设计", level=1)
    add_bullets(doc, [
        "密码不得明文存储，应使用加盐哈希。",
        "权限控制采用 RBAC 与管辖范围双重校验，前端声明不作为可信结论。",
        "学生只能查看本人预约和信用明细；管理员只能访问管辖范围内数据。",
        "扣分、封禁、角色变更、参数发布等敏感操作必须写操作日志。",
        "正式环境建议开启 HTTPS，数据库文件不对客户端直接开放。"
    ])

    doc.add_heading("9．维护设计", level=1)
    add_para(doc, "系统按模块化单体设计，业务能力通过 Store 和回调接口隔离。新增角色、扣分规则、恢复规则优先通过 RBAC 和 system_parameter 配置扩展；数据库访问集中在各模块 Store 中，便于后续将 SQLite 替换为 MySQL 或 PostgreSQL。日志、状态字典和测试用例用于后续定位问题和验证修改。")
    normalize_heading_runs(doc)
    doc.save(OVERVIEW_OUT)


def module_program_rows() -> list[list[str]]:
    return [
        ["A 平台与集成", "登录、角色切换、用户/角色/权限、管辖范围、系统参数、操作日志。", "用户凭证、session、角色、权限码、参数键值。", "当前用户上下文、权限校验结果、参数对象、操作日志。", "密码校验后建立 session；按 RBAC 读取权限；按 admin_lab_scope 校验数据范围；参数修改时版本递增并写 history。", "非法角色切换、越权访问、参数修改、日志查询。"],
        ["B 资源与预约", "提交预约、幂等重试、并发冲突检测、审批、驳回、取消、维护取消、签到状态联动。", "Idempotency-Key、lab_id、equipment_id、start_time、end_time、purpose、comment、maintenance_cancel。", "reservation 对象；late_cancel、credit_deduction_required；错误状态码。", "统一使用服务端回调取得身份/权限/信用/资源状态；预约创建在 BEGIN IMMEDIATE 事务中按半开区间查冲突；request_id 唯一保证幂等。", "20 线程抢约、缺回调 503、越权审批 403、临期取消扣分标记、统一入口挂载。"],
        ["C 现场与设备", "实验室设备、可用时段、资源可约判断、签到签退、IoT 指令、报修维修。", "设备编号、状态、时段规则、预约编号、二维码令牌、定位信息、报修描述。", "设备/时段对象、签到记录、IoT 日志、报修单、受影响预约处理结果。", "按设备状态、时段规则和预约占用判断可约性；签到通过后调用 B.mark_using；签退后调用 B.mark_completed；维修时调用 B 查询和取消受影响预约。", "设备维修不可约、二维码过期、定位失败人工核验、IoT 失败不回滚签到。"],
        ["D 信用与运营", "违规扫描、信用扣分、封禁、恢复申请、通知、统计。", "预约/签到/取消/报修记录、违规类型、恢复申请、统计时间范围。", "violation_record、credit_log、ban_record、notification、统计结果。", "定时任务扫描异常行为；管理员确认后写信用日志并判断封禁；恢复申请关联原扣分记录；通知使用唯一业务键避免重复。", "临期取消、爽约、迟到、重复扣分、信用恢复、权限范围统计。"],
    ]


def build_detail() -> None:
    doc = Document()
    apply_styles(doc)
    add_title_page(doc, "详细设计说明书")
    add_toc_placeholder(doc)

    doc.add_heading("1．引言", level=1)
    doc.add_heading("1.1 编写目的", level=2)
    add_para(doc, "本文档在概要设计基础上进一步说明各程序模块的功能、输入输出、算法、程序逻辑、接口、存储分配、限制条件和测试要点，为后续编码、联调、测试和文档整合提供依据。")
    doc.add_heading("1.2 项目背景", level=2)
    add_para(doc, "智能实验室预约与设备管理系统以在线 Web 应用方式解决实验室资源预约冲突、设备故障处理、信用追踪和统计分析问题。当前 A/B/C 已围绕统一 Flask 主入口形成集成基线，B 模块代码已完成本地测试。")
    doc.add_heading("1.3 定义", level=2)
    add_table(doc, ["术语", "说明"], [
        ["Store", "各模块封装数据库访问和业务规则的服务对象，如 PlatformStore、ReservationStore。"],
        ["回调注入", "统一入口把 A/C/D 的服务端能力放入 Flask config，B 在运行时调用。"],
        ["半开区间", "预约时间按 [start_time, end_time) 判断冲突，首尾相接不冲突。"],
        ["维护取消", "设备维修或停用导致的管理员取消，不触发学生信用扣分。"],
    ], [2000, 7400])
    doc.add_heading("1.4 参考资料", level=2)
    add_bullets(doc, [
        "《需求规格说明书0706.docx》",
        "《概要设计说明书_初稿_2026-07-11.docx》",
        "《A_统一Flask主入口与线上部署交接文档_2026-07-10.md》",
        "《B_资源与预约模块统一入口交接文档_2026-07-11.md》",
        "LC_proj/b_reservation/service.py、api.py、tests/test_*.py",
        "软件工程文档模板 05_详细设计说明书.dot",
    ])

    doc.add_heading("2．总体设计", level=1)
    doc.add_heading("2.1 需求概述", level=2)
    add_para(doc, "系统围绕学生预约使用实验室资源的全过程设计：登录鉴权、查询资源、提交预约、管理员审批、现场签到、使用签退、设备报修、违规信用处理、通知日志和统计分析。")
    doc.add_heading("2.2 软件结构", level=2)
    add_table(doc, ["层次", "组成", "说明"], [
        ["表现层", "浏览器页面、API 调用、后续演示页面", "负责表单输入、列表展示、状态提示和确认操作。"],
        ["应用入口层", "unified_app:create_app()", "创建 Flask 应用、读取配置、初始化 Store、注册路由、注入回调。"],
        ["业务模块层", "A/B/C/D Store 与 API", "按职责封装权限、预约、设备、信用、通知和统计逻辑。"],
        ["数据访问层", "sqlite3、schema.sql、短事务", "执行 SQL、事务控制、唯一约束和状态更新。"],
        ["基础设施层", "Gunicorn、Nginx、systemd、备份脚本", "提供线上运行、日志和数据库备份。"],
    ], [1500, 2800, 5100])
    add_para(doc, "B 模块公开 `create_app(config)` 用于独立测试，公开 `register_routes(app, store, register_health=False)` 用于统一入口挂载。统一入口接入时 B 不注册自己的 `/health`，由总入口统一返回模块列表。")

    doc.add_heading("3．程序描述", level=1)
    add_para(doc, "本节按 A/B/C/D 四个主要程序模块描述。B 模块已有码级实现，因此描述到类、方法和核心逻辑；A/C/D 根据现有交接资料和需求规格说明书描述到接口契约和实现约束。")

    for idx, (module, func, inputs, outputs, algo, tests) in enumerate(module_program_rows(), start=1):
        prefix = f"3.{idx}"
        doc.add_heading(f"{prefix} {module}", level=2)
        doc.add_heading(f"{prefix}.1 功能", level=3)
        add_para(doc, func)
        doc.add_heading(f"{prefix}.2 性能", level=3)
        if module.startswith("B"):
            add_para(doc, "预约无冲突情况下 95% 在 1.5 秒内完成事务；同一设备同一时段 20 个并发提交时，只允许一个有效预约成功，其余请求返回冲突。")
        else:
            add_para(doc, "遵循需求规格说明书性能目标：普通接口 95% 在 1~2 秒级返回，统计类接口普通时间范围 95% 在 5 秒内返回。")
        doc.add_heading(f"{prefix}.3 输入项目", level=3)
        add_para(doc, inputs)
        doc.add_heading(f"{prefix}.4 输出项目", level=3)
        add_para(doc, outputs)
        doc.add_heading(f"{prefix}.5 算法", level=3)
        add_para(doc, algo)
        doc.add_heading(f"{prefix}.6 程序逻辑", level=3)
        if module.startswith("B"):
            add_numbered(doc, [
                "HTTP 层调用 CURRENT_ACTOR()，若未登录或 user_id 非整数则返回权限错误。",
                "创建预约时读取 Idempotency-Key、lab_id、equipment_id、start_time、end_time、purpose。",
                "通过 STUDENT_IS_ELIGIBLE(user_id) 判断信用/封禁资格，通过 RESOURCE_IS_BOOKABLE(...) 判断资源可约性。",
                "ReservationStore.create() 标准化时间，校验正整数 ID、目的长度、结束时间晚于开始时间。",
                "使用 BEGIN IMMEDIATE 开启写事务；若 request_id 已存在，直接返回已有预约，实现幂等。",
                "按 active 状态 pending、approved、using 查询时间重叠；重叠条件为已有 start_time < 新 end_time 且已有 end_time > 新 start_time。",
                "无冲突时写入 pending 预约；审批、驳回、取消在事务中校验当前状态后更新。",
                "若 AUDIT 可用，创建、审批、驳回、取消均写操作日志。"
            ])
        elif module.startswith("A"):
            add_numbered(doc, [
                "登录成功后将用户上下文写入服务端 session。",
                "角色切换时校验当前用户是否绑定该角色，非法切换拒绝并审计。",
                "权限判断先检查角色权限，再检查实验室管辖范围。",
                "参数更新时写 system_parameter，同时记录 system_parameter_history 和 operation_log。",
            ])
        elif module.startswith("C"):
            add_numbered(doc, [
                "资源可约判断综合设备状态、可用时段规则、维护/停用规则和 B 的预约占用。",
                "签到先校验预约状态、预约人、时间窗口、二维码令牌和定位/人工核验。",
                "签到成功写 check_record 并调用 B.mark_using；IoT 失败只写 failed 日志，不回滚签到。",
                "报修后设备进入 fault_pending 或 repairing，维修影响已有预约时调用 B 的维护取消能力。"
            ])
        else:
            add_numbered(doc, [
                "定时任务扫描预约、签到、签退和取消记录，生成疑似违规。",
                "管理员确认违规后写 credit_log，更新分值并判断是否生成 ban_record。",
                "信用恢复申请必须关联原扣分记录，审核通过后生成恢复类 credit_log。",
                "通知根据业务事件生成，使用唯一业务键避免重复通知。"
            ])
        doc.add_heading(f"{prefix}.7 接口", level=3)
        if module.startswith("B"):
            add_table(doc, ["接口", "说明"], [
                ["create_app(config)", "独立创建 B 预约服务。"],
                ["register_routes(app, store, register_health=False)", "把 B 路由挂载到统一 Flask app。"],
                ["ReservationStore.create(...)", "创建预约，包含幂等、资格、资源可约和冲突判断。"],
                ["ReservationStore.approve/reject/cancel(...)", "审批、驳回和取消预约。"],
                ["mark_using/mark_completed", "供 C 签到签退后更新预约状态。"],
                ["list_affected_by_equipment/cancel_for_maintenance", "供 C 设备维护时处理受影响预约。"],
            ], [3400, 6000])
        else:
            add_para(doc, "接口详见概要设计内部接口表。模块对外暴露 HTTP API 或统一入口回调，模块之间不直接复制对方的权限、资源或信用规则。")
        doc.add_heading(f"{prefix}.8 存储分配", level=3)
        if module.startswith("B"):
            add_para(doc, "B 模块当前只直接维护 reservation 表，字段包括 reservation_id、request_id、user_id、lab_id、equipment_id、start_time、end_time、purpose、status、approver_id、approved_at、approve_comment、created_at。")
        else:
            add_para(doc, "模块数据表见概要设计 5.2。所有模块共用统一 SQLite 数据库文件，路径由 DATABASE_PATH 注入。")
        doc.add_heading(f"{prefix}.9 限制条件", level=3)
        if module.startswith("B"):
            add_para(doc, "B 不直接维护信用分、封禁、设备状态和通知落库；这些结论必须由 A/C/D 的服务端回调提供。当前 SQLite 使用数据库级短写锁，课程规模可接受，若迁移数据库再改为资源行锁。")
        else:
            add_para(doc, "课程阶段允许使用模拟接口和简化页面，但数据库字段、服务端校验、权限边界和日志追踪应保持一致。")
        doc.add_heading(f"{prefix}.10 测试要点", level=3)
        add_para(doc, tests)

    doc.add_heading("3.5 B 模块类与方法细化", level=2)
    add_table(doc, ["文件", "类/函数", "职责"], [
        ["b_reservation/api.py", "create_app(config)", "创建独立 Flask 应用，读取 DATABASE_PATH 和默认配置，初始化 ReservationStore。"],
        ["b_reservation/api.py", "register_routes(app, store, register_health)", "注册 B 模块 HTTP 路由，可挂载到统一入口。"],
        ["b_reservation/api.py", "actor()", "从 CURRENT_ACTOR 回调取得当前用户，未登录则抛 PermissionDenied。"],
        ["b_reservation/api.py", "can_manage(current, lab_id)", "调用 CAN_MANAGE_LAB 校验管理员管辖范围。"],
        ["b_reservation/api.py", "audit(...)", "若 A 注入 AUDIT，则记录预约创建、审批、驳回、取消日志。"],
        ["b_reservation/service.py", "ReservationStore.connect()", "创建 SQLite 连接，开启 foreign_keys 和 busy_timeout。"],
        ["b_reservation/service.py", "ReservationStore.create()", "预约创建、幂等、事务、冲突检测和写入。"],
        ["b_reservation/service.py", "ReservationStore.approve()/reject()", "审批通过或驳回待审批预约。"],
        ["b_reservation/service.py", "ReservationStore.cancel()", "学生取消或管理员维护取消，返回信用扣分标记。"],
        ["b_reservation/service.py", "mark_using()/mark_completed()", "C 签到签退后更新预约状态。"],
        ["b_reservation/service.py", "list_affected_by_equipment()", "设备维护时查询受影响有效预约。"],
        ["b_reservation/service.py", "student_has_equipment_relation()", "C 报修时判断学生近期是否使用过设备。"],
        ["b_reservation/service.py", "_find_conflict()", "按半开区间查找同设备或同实验室冲突预约。"],
    ], [2600, 2800, 4300])

    doc.add_heading("3.6 B 模块核心伪代码", level=2)
    pseudo = [
        "function create_reservation(request):",
        "    current = CURRENT_ACTOR(); require current.user_id",
        "    payload = JSON request body; request_id = Idempotency-Key",
        "    eligible = STUDENT_IS_ELIGIBLE(current.user_id)",
        "    bookable = RESOURCE_IS_BOOKABLE(lab_id, equipment_id, start_time, end_time)",
        "    begin immediate transaction",
        "        if request_id exists: return existing reservation",
        "        if time interval conflicts with pending/approved/using reservation: rollback and return conflict",
        "        insert reservation(status='pending')",
        "    commit transaction",
        "    AUDIT('reservation.create') if available",
        "    return reservation",
    ]
    for line in pseudo:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        run.font.size = Pt(9.5)

    doc.add_heading("3.7 测试设计摘要", level=2)
    add_table(doc, ["测试项", "输入/场景", "期望结果"], [
        ["幂等提交", "同一 Idempotency-Key 重复提交", "返回同一 reservation，不新增记录。"],
        ["时间冲突", "同一设备 09:00-10:00 与 09:30-10:30", "后提交请求返回冲突。"],
        ["首尾相接", "09:00-10:00 与 10:00-11:00", "允许预约。"],
        ["20 线程抢约", "20 个线程提交同一设备同一时段", "1 个成功，19 个冲突。"],
        ["越权审批", "非管辖管理员审批", "返回 403。"],
        ["缺少集成", "未注入 CURRENT_ACTOR 等回调", "业务接口返回 503。"],
        ["统一入口", "register_routes(..., register_health=False)", "不注册 B 自己的 /health，可由总入口接管。"],
        ["C 联动", "签到、签退、设备维护取消", "预约状态正确流转，维护取消不扣信用。"],
    ], [2300, 3300, 3800])
    normalize_heading_runs(doc)
    doc.save(DETAIL_OUT)


def write_artifact() -> None:
    ARTIFACT_OUT.write_text(
        f"""# 设计文档模板提取记录

- 日期：{TODAY}
- 模板来源：文档与绘图/软件工程文档模板.zip
- 封面参考：文档与绘图/组号-封面格式.docx
- 04 模板章节：引言、任务概述、总体设计、接口设计、数据结构设计、运行设计、出错处理设计、安全保密设计、维护设计。
- 05 模板章节：引言、总体设计、程序描述；程序描述包含功能、性能、输入项目、输出项目、算法、程序逻辑、接口、存储分配、限制条件、测试要点。
- 生成策略：保留模板章节体系，使用当前需求规格说明书、A/B 交接材料、B 代码与测试内容填充；A/C/D 写到整合接口粒度，B 写到类方法和核心逻辑粒度。
- 标题字体：一级、二级、三级标题的样式与具体 run 均显式统一为“微软雅黑”，避免 Word/WPS 对中文、数字或局部字符进行字体回退。
- QA 限制：当前环境未找到 LibreOffice，Word COM 因登录会话限制不可用，无法完成 DOCX->PNG 视觉渲染；已进行 python-docx 结构化生成和标题/表格/关键文本检查。
""",
        encoding="utf-8",
    )


def main() -> None:
    write_artifact()
    build_overview()
    build_detail()
    print(OVERVIEW_OUT)
    print(DETAIL_OUT)
    print(ARTIFACT_OUT)


if __name__ == "__main__":
    main()
