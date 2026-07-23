from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "文档与绘图" / "230710_8_需求规格说明书.docx"
OUTPUT = ROOT / "文档与绘图" / "230710_8_需求规格说明书_B首版技术基准补充_高亮.docx"


def highlight_run(run):
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return run


def set_paragraph(paragraph, text):
    paragraph.clear()
    highlight_run(paragraph.add_run(text))


def set_cell(cell, text):
    cell.text = ""
    highlight_run(cell.paragraphs[0].add_run(text))


def highlight_row(row):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                highlight_run(run)


def insert_after(paragraph, text):
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    created.style = paragraph.style
    highlight_run(created.add_run(text))
    return created


def find_paragraph(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.strip().startswith(prefix))


def add_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)
    table._tbl.tblPr.append(borders)


doc = Document(SOURCE)

# Revision identity.
set_paragraph(find_paragraph(doc, "文档版本："), "文档版本：V1.1")
set_paragraph(find_paragraph(doc, "完成日期："), "完成日期：2026 年 7 月 1 日")
revision = doc.tables[0].add_row()
values = [
    "V1.1",
    "2026-07-01",
    "依据 B 模块首个可运行版本冻结首版技术基准；补充 Python/Flask/SQLite、在线部署、时间与接口格式、预约表字段、幂等和并发实现约定",
    "230710_8 小组",
]
for cell, value in zip(revision.cells, values):
    set_cell(cell, value)

# 2.3.2: replace the old list of optional stacks with the implementation baseline.
environment = doc.tables[7]
set_cell(environment.cell(0, 1), "统一暂定基准（B 首版，2026-07-01）")
baseline_rows = [
    ("部署形态", "单服务器上的模块化单体 Web 系统；浏览器通过 HTTP/HTTPS 访问，后端与数据库均运行或存放在服务器端，客户端不得直连数据库"),
    ("语言与运行环境", "Python 3.12 及以上；开发机当前为 3.13.0，线上服务器为 3.12.3；组内开发使用 conda 的 csv 环境，服务器使用 venv 并按 requirements.txt 安装依赖"),
    ("Web 框架", "Flask 3.1.3；全组共用一个 Flask 应用，不为 A/B/C/D 分别建立独立后端"),
    ("数据库", "SQLite 3.45 及以上（Python sqlite3）；线上服务器为 3.45.1，开发机当前为 3.53.0；单一数据库文件由服务器持有，路径通过配置传入，不写个人绝对路径"),
    ("SQLite 类型映射", "实际建表统一使用 INTEGER、REAL、TEXT、BLOB；旧附录中的 bigint 映射为 INTEGER，varchar/text/datetime 映射为 TEXT，日期时间按本文件 ISO 8601 约定存储"),
    ("数据与接口编码", "UTF-8；HTTP 接口使用 JSON；成功或失败以 HTTP 状态码区分，错误体至少包含 error 字段"),
    ("时间约定", "服务器时区统一设为 Asia/Shanghai；接口和数据库使用 ISO 8601 本地时间 YYYY-MM-DDTHH:MM:SS，精确到秒，当前版本不接收带时区偏移的时间"),
    ("依赖与启动", "Python 依赖统一记录在 requirements.txt；数据库结构通过 schema.sql 或后续统一初始化入口创建"),
    ("缓存与异步", "首版不引入 Redis、消息队列或微服务；通知和定时任务采用同一后端内的简单实现"),
    ("IoT 接口", "课程范围内以模拟接口和日志记录实现"),
]
while len(environment.rows) < len(baseline_rows) + 1:
    environment.add_row()
for row, (name, value) in zip(environment.rows[1:], baseline_rows):
    set_cell(row.cells[0], name)
    set_cell(row.cells[1], value)

# 3.3: document the actual database/transaction baseline.
insert_after(
    find_paragraph(doc, "系统采用关系型数据库存储核心业务数据"),
    "【V1.1 首版实现基准】当前统一采用服务器端 SQLite。连接开启 foreign_keys，写操作设置 10 秒 busy timeout；预约创建、审批、驳回和取消均使用短事务。预约创建使用 BEGIN IMMEDIATE，使冲突检查和写入处于同一事务；若后续统一迁移至 MySQL 或 PostgreSQL，业务状态、字段语义和测试保持不变，仅替换连接层，并改用目标资源行锁。",
)

# Core table summary: make the reservation row match the shipped schema.
reservation_summary = doc.tables[13].rows[10]
set_cell(
    reservation_summary.cells[2],
    "reservation_id, request_id, user_id, lab_id, equipment_id, start_time, end_time, status",
)
set_cell(
    reservation_summary.cells[3],
    "B 模块拥有；存储预约、审批与取消结果，request_id 为唯一幂等键",
)

# FR-03: align prose and concurrency rules with the actual API and overlap predicate.
set_paragraph(
    find_paragraph(doc, "4. 系统生成请求幂等号"),
    "4. 客户端为一次业务提交生成幂等号，并通过 Idempotency-Key 请求头提交；后端将其写入 reservation.request_id 唯一字段。相同幂等号重试返回同一预约，不新增记录。",
)
set_paragraph(
    find_paragraph(doc, "5. 系统开启数据库事务"),
    "5. 系统开启数据库短事务。SQLite 首版使用 BEGIN IMMEDIATE 串行化写入；迁移到支持行锁的数据库时按目标资源加锁。",
)
concurrency = doc.tables[23]
extra_rules = [
    ("CON-FR03-05", "重叠判定采用半开区间：[start_time, end_time)。已有开始时间 < 新结束时间且已有结束时间 > 新开始时间时判定冲突；首尾相接不冲突"),
    ("CON-FR03-06", "request_id 长度为 1～64 个字符并具有唯一约束；同一幂等号重复提交不得新增预约"),
]
for code, rule in extra_rules:
    row = concurrency.add_row()
    set_cell(row.cells[0], code)
    set_cell(row.cells[1], rule)

# 6.3: specify the common HTTP trust boundary and B's exact first API surface.
insert_after(
    find_paragraph(doc, "6.3 软件接口"),
    "【V1.1 接口基准】浏览器仅提交业务数据，不得提交或决定“是否有资格预约”“资源是否可预约”“是否有管辖权”等可信结论；这些值必须由后端通过 A/C/D 模块能力取得。B 首版统一使用 JSON，请求成功返回 2xx，参数错误返回 400，权限不足返回 403，业务冲突或非法状态返回 409，依赖模块尚未接入返回 503。",
)
software_interfaces = doc.tables[47]
set_cell(software_interfaces.rows[2].cells[0], "创建预约 POST /api/reservations")
set_cell(software_interfaces.rows[2].cells[1], "请求头 Idempotency-Key；JSON：lab_id、equipment_id、start_time、end_time、purpose")
set_cell(software_interfaces.rows[2].cells[2], "reservation 对象；201")
set_cell(software_interfaces.rows[3].cells[0], "审批/驳回 POST /api/reservations/{id}/approve|reject")
set_cell(software_interfaces.rows[3].cells[1], "预约编号；JSON：comment；当前登录管理员由服务端会话取得")
set_cell(software_interfaces.rows[3].cells[2], "更新后的 reservation 对象；200")
for values in [
    ("个人预约 GET /api/reservations/me", "当前登录用户由服务端会话取得", "reservations 列表；200"),
    ("待审批 GET /api/labs/{lab_id}/reservations/pending", "实验室编号；服务端校验管辖范围", "reservations 列表；200 或 403"),
    ("取消预约 POST /api/reservations/{id}/cancel", "预约编号；维护取消时 JSON：maintenance_cancel=true", "更新后的 reservation，含 late_cancel 和 credit_deduction_required；200/403/409"),
]:
    row = software_interfaces.add_row()
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)

# Appendix: exact B-owned reservation table from schema.sql.
heading = doc.add_heading("B.9 reservation 预约表（B 首版基准）", level=2)
for run in heading.runs:
    highlight_run(run)
note = doc.add_paragraph()
highlight_run(
    note.add_run(
        "本表与 LC_proj/b_reservation/schema.sql 一致。当前独立版本在 A/C 共享表尚未合并前不声明物理外键；合并后由表负责人补充 user_id、lab_id、equipment_id 和 approver_id 的外键，不改变字段含义。"
    )
)
table = doc.add_table(rows=1, cols=4)
add_table_borders(table)
headers = ["字段名", "SQLite 类型", "约束", "说明"]
for cell, value in zip(table.rows[0].cells, headers):
    set_cell(cell, value)
reservation_fields = [
    ("reservation_id", "INTEGER", "PK, AUTOINCREMENT", "预约编号"),
    ("request_id", "TEXT", "unique, not null", "请求幂等号，1～64 个字符"),
    ("user_id", "INTEGER", "not null", "预约学生编号；合并后关联 user_account"),
    ("lab_id", "INTEGER", "not null", "实验室编号；合并后关联 laboratory"),
    ("equipment_id", "INTEGER", "nullable", "设备编号；为空表示预约实验室资源"),
    ("start_time", "TEXT", "not null", "预约开始时间，ISO 8601，本地时间精确到秒"),
    ("end_time", "TEXT", "not null", "预约结束时间，必须晚于开始时间"),
    ("purpose", "TEXT", "not null", "预约事由，1～255 个字符"),
    ("status", "TEXT", "not null, CHECK", "pending、approved、rejected、cancelled、using、completed、suspected_violation、no_show、violation_processed"),
    ("approver_id", "INTEGER", "nullable", "审批管理员编号"),
    ("approved_at", "TEXT", "nullable", "审批时间"),
    ("approve_comment", "TEXT", "nullable", "审批意见，最多 255 个字符"),
    ("created_at", "TEXT", "not null", "创建时间"),
]
for values in reservation_fields:
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)

doc.core_properties.title = "智能实验室预约与设备管理系统需求规格说明书 V1.1"
doc.save(OUTPUT)
print(OUTPUT)
