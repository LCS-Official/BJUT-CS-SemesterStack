from __future__ import annotations

import json
import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from lxml import etree


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def rgb_text(color) -> str | None:
    if color is None or color.rgb is None:
        return None
    return str(color.rgb)


def inspect(path: Path) -> dict:
    doc = Document(path)
    paragraphs = []
    highlights = []
    styles = Counter()
    fonts = Counter()

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        styles[paragraph.style.name if paragraph.style else ""] += 1
        run_info = []
        num_id = None
        ilvl = None
        ppr = paragraph._p.pPr
        if ppr is not None and ppr.numPr is not None:
            if ppr.numPr.numId is not None:
                num_id = ppr.numPr.numId.val
            if ppr.numPr.ilvl is not None:
                ilvl = ppr.numPr.ilvl.val
        for run in paragraph.runs:
            if not run.text:
                continue
            font_key = (
                run.font.name,
                run.font.size.pt if run.font.size else None,
                bool(run.bold),
                rgb_text(run.font.color),
            )
            fonts[str(font_key)] += len(run.text)
            highlight = run.font.highlight_color
            if highlight is not None:
                item = {
                    "paragraph_index": i,
                    "paragraph": text,
                    "run": run.text,
                    "highlight": str(highlight),
                    "yellow": highlight == WD_COLOR_INDEX.YELLOW,
                }
                highlights.append(item)
            run_info.append(
                {
                    "text": run.text,
                    "font": run.font.name,
                    "size_pt": run.font.size.pt if run.font.size else None,
                    "bold": bool(run.bold),
                    "italic": bool(run.italic),
                    "highlight": str(highlight) if highlight is not None else None,
                }
            )
        paragraphs.append(
            {
                "index": i,
                "style": paragraph.style.name if paragraph.style else None,
                "text": text,
                "num_id": num_id,
                "ilvl": ilvl,
                "runs": run_info,
            }
        )

    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"index": ti, "rows": rows})

    style_definitions = {}
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "toc 1", "toc 2", "Header", "Footer"):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        pf = style.paragraph_format
        style_definitions[style_name] = {
            "font": style.font.name,
            "size_pt": style.font.size.pt if style.font.size else None,
            "bold": style.font.bold,
            "italic": style.font.italic,
            "color": rgb_text(style.font.color),
            "alignment": str(pf.alignment) if pf.alignment is not None else None,
            "left_indent_pt": pf.left_indent.pt if pf.left_indent else None,
            "right_indent_pt": pf.right_indent.pt if pf.right_indent else None,
            "first_line_indent_pt": pf.first_line_indent.pt if pf.first_line_indent else None,
            "space_before_pt": pf.space_before.pt if pf.space_before else None,
            "space_after_pt": pf.space_after.pt if pf.space_after else None,
            "line_spacing": float(pf.line_spacing) if isinstance(pf.line_spacing, (int, float)) else (pf.line_spacing.pt if pf.line_spacing else None),
            "keep_with_next": pf.keep_with_next,
            "page_break_before": pf.page_break_before,
        }

    xml_marks = []
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        images = [n for n in names if n.startswith("word/media/")]
        headers_footers = {}
        for name in names:
            if name.startswith("word/header") or name.startswith("word/footer"):
                root = etree.fromstring(archive.read(name))
                headers_footers[name] = "".join(root.itertext()).strip()
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            try:
                root = etree.fromstring(archive.read(name))
            except etree.XMLSyntaxError:
                continue
            for paragraph in root.iter(W + "p"):
                paragraph_text = "".join(paragraph.itertext()).strip()
                for run in paragraph.iter(W + "r"):
                    run_text = "".join(run.itertext())
                    rpr = run.find(W + "rPr")
                    if rpr is None:
                        continue
                    highlight = rpr.find(W + "highlight")
                    shading = rpr.find(W + "shd")
                    mark = None
                    if highlight is not None:
                        mark = {"kind": "highlight", "value": highlight.get(W + "val")}
                    elif shading is not None:
                        fill = shading.get(W + "fill")
                        if fill and fill.upper() not in {"AUTO", "FFFFFF", "000000"}:
                            mark = {"kind": "shading", "value": fill}
                    if mark:
                        xml_marks.append(
                            {
                                "part": name,
                                "paragraph": paragraph_text,
                                "run": run_text,
                                **mark,
                            }
                        )

    return {
        "path": str(path),
        "sections": len(doc.sections),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "image_part_count": len(images),
        "styles": styles.most_common(),
        "style_definitions": style_definitions,
        "fonts_weighted_chars": fonts.most_common(30),
        "highlights": highlights,
        "xml_marks": xml_marks,
        "paragraphs": paragraphs,
        "tables": tables,
        "headers_footers": headers_footers,
    }


def main() -> None:
    if len(sys.argv) < 3:
        raise SystemExit("usage: inspect_docx.py OUTPUT.json INPUT.docx [INPUT.docx ...]")
    output = Path(sys.argv[1])
    data = [inspect(Path(p)) for p in sys.argv[2:]]
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    for item in data:
        print(item["path"])
        print(
            f"  paragraphs={item['paragraph_count']} tables={item['table_count']} "
            f"images={item['image_part_count']} highlights={len(item['highlights'])} "
            f"xml_marks={len(item['xml_marks'])}"
        )
        for h in item["highlights"]:
            print(f"  HIGHLIGHT p{h['paragraph_index']}: {h['run']!r} | {h['paragraph']}")
        for mark in item["xml_marks"]:
            print(
                f"  XML_MARK {mark['part']} {mark['kind']}={mark['value']}: "
                f"{mark['run']!r} | {mark['paragraph']}"
            )


if __name__ == "__main__":
    main()
