from __future__ import annotations

import copy
import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


ROOT = Path(r"C:\Users\LC\Desktop\SoftWare_CD")
TEST_SOURCE = ROOT / "最终文档等" / "测试计划.docx"
TEST_TEMPLATE = ROOT / "tools" / "final_docs" / "07_测试计划_模板参考.docx"
TEST_FINAL = ROOT / "最终文档等" / "测试计划_最终版.docx"
REQ_SOURCE = ROOT / "文档与绘图" / "需求规格说明书0706.docx"
REQ_FINAL = ROOT / "最终文档等" / "需求规格说明书_最终版.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = "{" + W_NS + "}"


HUMAN_REWRITES = {
    "本计划规定智能实验室预约与设备管理系统的测试范围、方法、环境、人员、进度、用例和通过准则，用于系统性发现功能、权限、状态、并发、数据一致性和部署问题，并为最终验收与测试分析报告提供依据。":
        "本计划说明智能实验室预约与设备管理系统要测什么、怎样测试、由谁完成以及如何判断通过。测试结果用于记录问题、安排回归，并作为最终检查和测试分析报告的依据。",
    "被测系统为部署在云服务器上的模块化单体 Web 应用，包含平台权限、资源预约、现场设备、信用运营和整合适配层。用户通过浏览器访问，后端使用 Flask，数据保存在 SQLite。":
        "本系统部署在云服务器上，用户通过浏览器访问。后端采用统一的 Flask 入口，各功能模块按目录分开，主要包括平台权限、资源预约、现场设备和信用运营，数据保存在 SQLite 数据库中。",
    "验证 FR-01～FR-13 功能、角色权限和异常分支符合需求。":
        "检查 FR-01～FR-13 的主要功能、角色权限和异常处理是否符合需求。",
    "验证预约、设备、报修、信用和通知跨模块闭环的一致性。":
        "串联预约、设备、报修、信用和通知流程，检查跨模块数据是否一致。",
    "验证重复提交、并发抢约、后台扫描和通知生成具备幂等性。":
        "重复执行提交、并发抢约、后台扫描和通知任务，确认不会产生重复业务记录。",
    "验证 HTTPS、Session、密码、范围控制、审计和备份恢复。":
        "检查 HTTPS、Session、密码处理、数据范围、审计和备份恢复。",
    "验证课程规模下的响应性能、100 个在线会话和可重复演示能力。":
        "在课程设计规模下测试常用接口和 100 个在线会话，并确认演示流程可以重复执行。",
    "测试覆盖登录与 RBAC、资源查询、预约与审批取消、签到签退与人工核验、设备报修和时段、违规信用与恢复、基础数据和参数、通知日志、统计导出，以及安全、性能、备份和恢复。":
        "测试内容包括登录与权限、资源查询、预约及审批取消、签到签退与人工核验、设备报修和可用时段、违规信用与恢复、基础数据和参数、通知日志、统计导出，以及安全、性能、备份和恢复。",
    "采用需求驱动和风险驱动结合的方法。服务层和规则采用白盒单元测试；接口、权限、状态和事务采用 Flask 测试客户端；跨模块闭环采用集成测试；真实浏览器用于角色页面、定位权限和响应式检查；公网环境执行冒烟、性能、HTTPS 和恢复验证。":
        "测试按需求和风险高低安排。服务层规则使用单元测试，接口、权限、状态和事务使用 Flask 测试客户端，跨模块流程使用集成测试。角色页面、定位授权和不同屏幕宽度在真实浏览器中检查，公网环境再完成冒烟、性能、HTTPS 和恢复测试。",
    "测试顺序为：静态检查与初始化→单元测试→模块接口测试→四模块集成→角色系统测试→并发性能与安全→缺陷回归→验收彩排。":
        "测试先从静态检查和数据库初始化开始，再依次进行单元测试、模块接口测试、四模块集成、角色系统测试、并发与安全测试，最后完成缺陷回归和验收彩排。",
    "测试项目以需求编号为主线，每个项目同时检查正常流程、非法参数、权限、状态、数据落库、通知审计和重复执行。高风险项目包括同资源并发抢约、维修批量取消、信用扣分封禁、人工核验补签和系统管理员基础数据维护。":
        "测试项目按照需求编号编排。除正常流程外，还要检查非法参数、权限、状态变化、数据写入、通知与日志，以及同一操作重复执行的结果。重点检查同一资源并发抢约、维修影响预约、信用扣分封禁、人工核验补签和基础数据维护。",
    "每条用例记录版本和数据基线，执行准备、操作、观察页面与接口、查询数据库或日志、比较预期、保存证据、恢复数据。失败用例建立缺陷记录，修复后执行原用例和相关回归。":
        "执行每条用例前先记录版本和初始数据。操作后检查页面、接口、数据库和日志是否符合预期，并保存结果；如用例失败，则登记问题，修复后重跑原用例及相关用例，最后恢复测试数据。",
    "评价范围包括所有功能需求、三类角色、模块间接口、主要数据表和状态机、正常与异常流程、并发一致性、性能、安全、兼容、部署、备份恢复和文档一致性。真实硬件控制、多节点高可用、对象存储和恶意文件扫描不在课程验收范围。":
        "评价范围包括全部功能需求、三类角色、模块接口、主要数据表和状态变化，同时检查正常与异常流程、并发一致性、性能、安全、兼容、部署、备份恢复和文档对应关系。课程验收不包含真实硬件控制、多节点高可用、对象存储和恶意文件扫描。",
}


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def set_visible_text(paragraph, text: str) -> None:
    text_nodes = paragraph._p.findall(".//" + W + "t")
    if not text_nodes:
        paragraph.add_run(text)
        return
    text_nodes[0].text = text
    if text.startswith(" ") or text.endswith(" "):
        text_nodes[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    for node in text_nodes[1:]:
        node.text = ""


def replace_properties(element, old_properties, new_properties) -> None:
    if old_properties is not None:
        element.remove(old_properties)
    if new_properties is not None:
        element.insert(0, copy.deepcopy(new_properties))


def keep_only_paragraph_style(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for child in list(ppr):
        if child.tag != qn("w:pStyle"):
            ppr.remove(child)


def clear_run_formatting(paragraph) -> None:
    for run in paragraph.runs:
        rpr = run._r.rPr
        if rpr is not None:
            run._r.remove(rpr)


def next_id(parent, tag: str, attr: str) -> int:
    values = []
    for element in parent.findall(qn(tag)):
        value = element.get(qn(attr))
        if value is not None:
            values.append(int(value))
    return max(values, default=0) + 1


def add_numbering(doc: Document, paragraphs, unique_id: int) -> None:
    numbering = doc.part.numbering_part.element
    abstract_id = unique_id
    num_id = unique_id

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{unique_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    level.append(text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    level.append(suff)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "420")
    ind.set(qn("w:hanging"), "315")
    ppr.append(ind)
    level.append(ppr)
    abstract.append(level)
    first_num_index = next(
        (i for i, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)

    for paragraph in paragraphs:
        ppr = paragraph._p.get_or_add_pPr()
        old = ppr.find(qn("w:numPr"))
        if old is not None:
            ppr.remove(old)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_el])
        ppr.append(num_pr)


def paragraphs_between(doc: Document, start_text: str, end_text: str) -> list:
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == start_text)
    end = next(i for i, p in enumerate(paragraphs[start + 1 :], start + 1) if p.text.strip() == end_text)
    return [p for p in paragraphs[start + 1 : end] if p.text.strip() and p.style.name == "Normal"]


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + tag))
        if node is None:
            node = OxmlElement("w:" + tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_tables(doc: Document) -> None:
    for table in doc.tables:
        table.autofit = True
        tbl_pr = table._tbl.tblPr
        shading = tbl_pr.find(qn("w:shd"))
        if shading is not None:
            tbl_pr.remove(shading)
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge = borders.find(qn("w:" + side))
            if edge is None:
                edge = OxmlElement("w:" + side)
                borders.append(edge)
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "4")
            edge.set(qn("w:color"), "000000")

        if table.rows:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:tblHeader")) is None:
                repeat = OxmlElement("w:tblHeader")
                repeat.set(qn("w:val"), "true")
                tr_pr.append(repeat)

        for row_index, row in enumerate(table.rows):
            row.height = None
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = tc_pr.find(qn("w:shd"))
                if shd is not None:
                    shd.set(qn("w:fill"), "FFFFFF")
                    shd.set(qn("w:val"), "clear")
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles["Normal"]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = None
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.bold = row_index == 0
                        run.font.name = "宋体"
                        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
                        run.font.size = Pt(10.5)


def replace_package_parts(target: Path, reference: Path, parts: tuple[str, ...]) -> None:
    with zipfile.ZipFile(reference) as ref_zip:
        replacements = {name: ref_zip.read(name) for name in parts if name in ref_zip.namelist()}
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=target.parent) as temp:
        temp_path = Path(temp.name)
    try:
        with zipfile.ZipFile(target) as source_zip, zipfile.ZipFile(temp_path, "w") as output_zip:
            for item in source_zip.infolist():
                data = replacements.get(item.filename, source_zip.read(item.filename))
                output_zip.writestr(item, data)
        temp_path.replace(target)
    finally:
        temp_path.unlink(missing_ok=True)


def build_test_plan() -> None:
    shutil.copy2(TEST_SOURCE, TEST_FINAL)
    doc = Document(TEST_FINAL)
    template = Document(TEST_TEMPLATE)

    body_start = next(i for i, p in enumerate(doc.paragraphs) if p.style.name == "Heading 1" and p.text.strip() == "1．引言")
    front = doc.paragraphs[:body_start]
    title = front[0]

    for paragraph in list(front[1:]):
        if paragraph.style.name not in {"toc 1", "toc 2"}:
            remove_paragraph(paragraph)

    set_visible_text(title, "七、测试计划")
    replace_properties(title._p, title._p.pPr, template.paragraphs[0]._p.pPr)
    clear_run_formatting(title)
    if title.runs and template.paragraphs[0].runs:
        replace_properties(title.runs[0]._r, title.runs[0]._r.rPr, template.paragraphs[0].runs[0]._r.rPr)

    title._p.addnext(copy.deepcopy(template.paragraphs[1]._p))

    # 模板目录之后保留两行空白，再进入正文。
    last_toc = [p for p in doc.paragraphs if p.style.name in {"toc 1", "toc 2"}][-1]
    last_toc._p.addnext(copy.deepcopy(template.paragraphs[27]._p))
    last_toc._p.addnext(copy.deepcopy(template.paragraphs[26]._p))

    body_start = next(i for i, p in enumerate(doc.paragraphs) if p.style.name == "Heading 1" and p.text.strip() == "1．引言")
    for paragraph in doc.paragraphs[body_start:]:
        original = paragraph.text.strip()
        if original in HUMAN_REWRITES:
            set_visible_text(paragraph, HUMAN_REWRITES[original])
        if paragraph.style.name.startswith("Heading"):
            normalized = re.sub(r"^(\d+(?:\.\d+)+)\s+", r"\1", paragraph.text.strip())
            if normalized != paragraph.text.strip():
                set_visible_text(paragraph, normalized)
        keep_only_paragraph_style(paragraph)
        clear_run_formatting(paragraph)

    numbered_sections = (
        ("1.4参考资料", "2．任务概述"),
        ("2.1目标", "2.2运行环境"),
        ("2.4条件与限制", "3．计划"),
        ("3.3测试准备", "3.4测试机构及人员"),
        ("4.2.4允许偏差", "4.3进度"),
        ("4.5测试资料", "5．评价"),
    )
    for sequence_index, (start, end) in enumerate(numbered_sections, 100):
        items = paragraphs_between(doc, start, end)
        if items:
            add_numbering(doc, items, sequence_index)

    toc_ppr = {
        "toc 1": template.paragraphs[2]._p.pPr,
        "toc 2": template.paragraphs[3]._p.pPr,
    }
    for paragraph in doc.paragraphs:
        if paragraph.style.name in {"toc 1", "toc 2"}:
            replace_properties(paragraph._p, paragraph._p.pPr, toc_ppr[paragraph.style.name])
            clear_run_formatting(paragraph)

    format_tables(doc)
    doc.save(TEST_FINAL)

    replace_package_parts(
        TEST_FINAL,
        TEST_TEMPLATE,
        (
            "word/styles.xml",
            "word/fontTable.xml",
            "word/theme/theme1.xml",
            "word/header1.xml",
            "word/footer1.xml",
        ),
    )


def remove_yellow_highlights(source: Path, destination: Path) -> int:
    count = 0
    with zipfile.ZipFile(source) as source_zip, zipfile.ZipFile(destination, "w") as output_zip:
        for item in source_zip.infolist():
            data = source_zip.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                except etree.XMLSyntaxError:
                    pass
                else:
                    changed = False
                    for highlight in root.findall(".//" + W + "highlight"):
                        if (highlight.get(W + "val") or "").lower() == "yellow":
                            highlight.getparent().remove(highlight)
                            count += 1
                            changed = True
                    if changed:
                        data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            output_zip.writestr(item, data)
    return count


def main() -> None:
    build_test_plan()
    removed = remove_yellow_highlights(REQ_SOURCE, REQ_FINAL)
    print(f"测试计划：{TEST_FINAL}")
    print(f"需求规格说明书：{REQ_FINAL}")
    print(f"已取消黄色标记：{removed} 处")


if __name__ == "__main__":
    main()
