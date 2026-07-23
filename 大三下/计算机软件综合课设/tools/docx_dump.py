import sys

from docx import Document


sys.stdout.reconfigure(encoding="utf-8")
doc = Document(sys.argv[1])
p_start = int(sys.argv[2]) if len(sys.argv) > 2 else 0
p_end = int(sys.argv[3]) if len(sys.argv) > 3 else len(doc.paragraphs)
t_start = int(sys.argv[4]) if len(sys.argv) > 4 else 0
t_end = int(sys.argv[5]) if len(sys.argv) > 5 else len(doc.tables)

for index, paragraph in enumerate(doc.paragraphs[p_start:p_end], p_start):
    text = paragraph.text.strip()
    if text:
        print(f"P{index:04d} [{paragraph.style.name}] {text}")

for table_index, table in enumerate(doc.tables[t_start:t_end], t_start):
    print(f"\nTABLE {table_index}")
    for row_index, row in enumerate(table.rows):
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        print(f"R{row_index:03d}: " + " | ".join(cells))
